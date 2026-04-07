from __future__ import annotations

from pathlib import Path

from docx import Document

from secknow.text_processing.loaders.router import load_document_text


def test_load_text_and_markdown(tmp_path: Path) -> None:
    """txt/md 文件应能正确读取。"""
    txt = tmp_path / "sample.txt"
    md = tmp_path / "sample.md"
    txt.write_text("hello txt", encoding="utf-8")
    md.write_text("# 标题\n正文", encoding="utf-8")

    assert load_document_text(txt) == "hello txt"
    assert "标题" in load_document_text(md)


def test_load_docx(tmp_path: Path) -> None:
    """docx 文件应能提取段落文本。"""
    file_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("第二段")
    doc.save(file_path)

    text = load_document_text(file_path)
    assert "第一段" in text
    assert "第二段" in text


def test_load_pdf_with_monkeypatch(tmp_path: Path, monkeypatch) -> None:
    """pdf 路由应经 PyMuPDF 打开并由 pdf_advanced 处理各页。"""
    monkeypatch.setenv("PDF_OCR_ENABLED", "0")
    file_path = tmp_path / "sample.pdf"
    file_path.write_bytes(b"%PDF-1.4")

    class FakePage:
        def __init__(self, content: str):
            self._content = content

        def get_text(self, mode: str = "text") -> str:
            return self._content

        def get_pixmap(self, **_kwargs):
            raise AssertionError("单测不应触发 OCR 光栅化")

        def find_tables(self):
            class _TF:
                tables: list = []

            return _TF()

    class FakeDoc:
        def __init__(self) -> None:
            self._pages = [FakePage("第一页"), FakePage("第二页")]

        def __len__(self) -> int:
            return len(self._pages)

        def __getitem__(self, i: int) -> FakePage:
            return self._pages[i]

        def close(self) -> None:
            return None

    monkeypatch.setattr(
        "secknow.text_processing.loaders.pdf_loader.fitz.open",
        lambda *_a, **_kw: FakeDoc(),
    )

    text = load_document_text(file_path)
    assert "第一页" in text
    assert "第二页" in text
